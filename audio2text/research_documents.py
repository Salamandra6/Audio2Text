from __future__ import annotations

import html
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from .exporters import TranscriptionResult
from .research_analysis import build_sections


def unique_path(folder: Path, stem: str, suffix: str) -> Path:
    candidate = folder / f"{stem}{suffix}"
    counter = 2
    while candidate.exists():
        candidate = folder / f"{stem} ({counter}){suffix}"
        counter += 1
    return candidate


def plain_text(result: TranscriptionResult, sections: dict) -> str:
    lines = ["TRANSCRIPTOR DE AUDIO A TEXTO", Path(result.source_file).name, ""]
    if sections["metadata"]:
        lines += ["METADATOS", "----------"]
        lines += [f"{label}: {value}" for label, value in sections["metadata"]]
        lines.append("")
    if sections["summary"]:
        lines += ["RESUMEN EXTRACTIVO", "-------------------"]
        lines += [f"- {sentence}" for sentence in sections["summary"]]
        lines.append("")
    if sections["topics"]:
        lines += ["TEMAS PRINCIPALES SUGERIDOS", "---------------------------"]
        lines += [f"- {topic}" for topic in sections["topics"]]
        lines.append("")
    title = "TRANSCRIPCIÓN LIMPIA" if sections["mode"] == "clean" else "TRANSCRIPCIÓN LITERAL"
    if sections["timestamps"]:
        title += " CON MARCAS DE TIEMPO"
    lines += [title, "-" * len(title)]
    lines += sections["transcript"] or ["No se detectó voz transcribible."]
    return "\n".join(lines).rstrip() + "\n"


def write_docx(path: Path, result: TranscriptionResult, sections: dict) -> None:
    document = Document()
    document.styles["Normal"].font.name = "Aptos"
    document.styles["Normal"].font.size = Pt(10.5)
    title = document.add_heading("Transcriptor de Audio a Texto", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(Path(result.source_file).name)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if sections["metadata"]:
        document.add_heading("Metadatos", level=1)
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in sections["metadata"]:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
    if sections["summary"]:
        document.add_heading("Resumen extractivo", level=1)
        for sentence in sections["summary"]:
            document.add_paragraph(sentence, style="List Bullet")
    if sections["topics"]:
        document.add_heading("Temas principales sugeridos", level=1)
        for topic in sections["topics"]:
            document.add_paragraph(topic, style="List Bullet")

    heading = "Transcripción limpia" if sections["mode"] == "clean" else "Transcripción literal"
    if sections["timestamps"]:
        heading += " con marcas de tiempo"
    document.add_heading(heading, level=1)
    for line in sections["transcript"] or ["No se detectó voz transcribible."]:
        document.add_paragraph(line)
    notice = document.add_paragraph("Uso exclusivo de su destinatario. Prohibida su comercialización.")
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(path)


def write_pdf(path: Path, result: TranscriptionResult, sections: dict) -> None:
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="ResearchBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=13, spaceAfter=5))
    styles.add(ParagraphStyle(name="ResearchSmall", parent=styles["BodyText"], fontName="Helvetica", fontSize=8, leading=10, textColor=colors.HexColor("#4c5d70")))
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm, topMargin=18 * mm, bottomMargin=18 * mm, title="Transcriptor de Audio a Texto", author="Audio2Text")
    story = [Paragraph("Transcriptor de Audio a Texto", styles["Title"]), Paragraph(html.escape(Path(result.source_file).name), styles["Heading2"]), Spacer(1, 5 * mm)]

    if sections["metadata"]:
        story.append(Paragraph("Metadatos", styles["Heading2"]))
        data = [[Paragraph(f"<b>{html.escape(label)}</b>", styles["ResearchBody"]), Paragraph(html.escape(value), styles["ResearchBody"])] for label, value in sections["metadata"]]
        table = Table(data, colWidths=[48 * mm, 112 * mm])
        table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#9aabba")),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eef3f8")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ]))
        story += [table, Spacer(1, 5 * mm)]
    if sections["summary"]:
        story.append(Paragraph("Resumen extractivo", styles["Heading2"]))
        story += [Paragraph(f"- {html.escape(sentence)}", styles["ResearchBody"]) for sentence in sections["summary"]]
    if sections["topics"]:
        story.append(Paragraph("Temas principales sugeridos", styles["Heading2"]))
        story += [Paragraph(f"- {html.escape(topic)}", styles["ResearchBody"]) for topic in sections["topics"]]

    heading = "Transcripción limpia" if sections["mode"] == "clean" else "Transcripción literal"
    if sections["timestamps"]:
        heading += " con marcas de tiempo"
    story.append(Paragraph(heading, styles["Heading2"]))
    for line in sections["transcript"] or ["No se detectó voz transcribible."]:
        story.append(Paragraph(html.escape(line).replace("\n", "<br/>"), styles["ResearchBody"]))
    story += [Spacer(1, 7 * mm), Paragraph("Uso exclusivo de su destinatario. Prohibida su comercialización.", styles["ResearchSmall"])]
    doc.build(story)


def export_research_result(result: TranscriptionResult, output_dir: str | Path, formats: Iterable[str], content_options: dict | None = None) -> list[Path]:
    folder = Path(output_dir)
    folder.mkdir(parents=True, exist_ok=True)
    stem = Path(result.source_file).stem
    sections = build_sections(result, content_options)
    written = []
    for selected in formats:
        format_name = selected.lower().strip()
        if format_name == "txt":
            path = unique_path(folder, stem, ".txt")
            path.write_text(plain_text(result, sections), encoding="utf-8")
        elif format_name == "docx":
            path = unique_path(folder, stem, ".docx")
            write_docx(path, result, sections)
        elif format_name == "pdf":
            path = unique_path(folder, stem, ".pdf")
            write_pdf(path, result, sections)
        else:
            raise ValueError(f"Formato de documento no admitido: {selected}")
        written.append(path)
    return written
